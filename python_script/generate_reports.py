import io
import os
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Inches, Pt

INPUT_DIR = "CSV"
OUTPUT_DIR = "output-docs"
ASSETS_DIR = "assets"
HEADER_IMAGE = os.path.join(ASSETS_DIR, "Picture.png")

TEXT_SCHOOL = "PODAR WORLD SCHOOL, BADWAI BHOPAL"
TEXT_REPORT = "RESULT ANALYSIS 2025-26"


def parse_float(val):
    try:
        return float(val)
    except:
        return 0.0


def normalize_score(obtained_str, total_str):
    """Normalize score to be out of 80."""
    try:
        obt = float(obtained_str)
    except:
        obt = 0.0  # Handle 'AB', 'absent', empty, etc.

    try:
        tot = float(total_str)
    except:
        tot = 80.0  # Default if missing, though unlikely in a valid CSV

    if tot == 0:
        return 0.0

    return (obt / tot) * 80


def parse_csv(file_path):
    with open(file_path, "r", encoding="utf-8", errors="replace") as f:
        lines = f.readlines()

    students = []
    subjects = []

    # 1. Detect Subjects from Header
    # Look for the line containing "S. No.,Student Name,EXAM"
    header_idx = -1
    for i, line in enumerate(lines):
        parts = [p.strip() for p in line.split(",")]
        if len(parts) > 3 and parts[0] == "S. No." and parts[1] == "Student Name" and parts[2] == "EXAM":
            header_idx = i
            # Parse subjects: (Name, OutOff) pairs starting at index 3
            # Col 3: Subj1, Col 4: OutOff
            # Col 5: Subj2, Col 6: OutOff
            # ...
            idx = 3
            while idx < len(parts):
                subj_name = parts[idx]
                # Break if empty or explicitly "Out Off" (which shouldn't happen at start of pair)
                # But sometimes csv has trailing commas
                if not subj_name:
                    break
                subjects.append(subj_name)
                idx += 2  # Skip the "Out Off" column
            break

    if header_idx == -1:
        print(f"Could not find header row in {file_path}")
        return [], []

    print(f"Detected Subjects ({len(subjects)}): {subjects}")

    # 2. Parse Students
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        parts = [p.strip() for p in line.split(",")]

        # Check if start of student (Col 0 is digit)
        if len(parts) > 1 and parts[0].isdigit():
            # This is a student start
            s_no = parts[0]
            name = parts[1]

            student_exams = []
            # We assume 3 rows per student: PT I, TERM I, PT II
            # This block starts at `i`

            rows_indices = [i, i + 1, i + 2]
            for r_idx_offset, r_idx in enumerate(rows_indices):
                if r_idx >= len(lines):
                    break

                row_line = lines[r_idx].strip()
                row_parts = [p.strip() for p in row_line.split(",")]

                # Pad row_parts if needed
                while len(row_parts) < 3 + len(subjects) * 2:
                    row_parts.append("")

                exam_name = row_parts[2]
                if not exam_name:
                    # Fallback names based on offset order if missing in CSV
                    if r_idx_offset == 0:
                        exam_name = "PT I"
                    elif r_idx_offset == 1:
                        exam_name = "TERM I"
                    elif r_idx_offset == 2:
                        exam_name = "PT II"

                # Extract scores
                scores = {}
                col = 3
                for subj in subjects:
                    val = row_parts[col]
                    tot = row_parts[col + 1]
                    norm_val = normalize_score(val, tot)
                    scores[subj] = norm_val
                    col += 2

                student_exams.append({"exam": exam_name, "scores": scores})

            students.append({"s_no": s_no, "name": name, "exams": student_exams})

            i += 3  # Skip the 3 rows
            continue

        i += 1

    return students, subjects


def create_graph(student_data, subjects):
    exams = [e["exam"] for e in student_data["exams"]]

    # GROUPING:
    # X Axis: Subjects
    # Series: Exams

    # Prepare data: marks_per_exam[exam_index] = [score_subj1, score_subj2...]
    exam_series = []
    for e in student_data["exams"]:
        scores = []
        for subj in subjects:
            scores.append(e["scores"].get(subj, 0.0))
        exam_series.append(scores)

    x = np.arange(len(subjects))
    num_exams = len(exams)

    # Bar Sizing
    group_width = 0.8
    bar_width = group_width / num_exams

    # Colors for Exams
    colors = ["#4285F4", "#EA4335", "#34A853", "#FBBC05"]

    fig, ax = plt.subplots(figsize=(14, 7))

    # Add faded background rectangles for groups (Subjects)
    rect_width = group_width + 0.1
    for pos in x:
        left = pos - rect_width / 2
        right = pos + rect_width / 2
        ax.axvspan(left, right, color="#F0F0F0", alpha=0.5, zorder=0)

    # Gridlines
    ax.grid(axis="y", linestyle="--", alpha=0.6, zorder=1)

    # Remove top and right spines
    ax.spines["top"].set_visible(False)
    ax.spines["right"].set_visible(False)

    for i, exam_name in enumerate(exams):
        marks = exam_series[i]

        total_group_width = num_exams * bar_width
        start_shift = -total_group_width / 2 + bar_width / 2
        offset = start_shift + i * bar_width

        ax.bar(
            x + offset,
            marks,
            bar_width,
            label=exam_name,
            color=colors[i % len(colors)],
            zorder=3,
            edgecolor="white",
            linewidth=0.5,
        )

    ax.set_ylabel("Marks (out of 80)", fontsize=12, fontweight="bold")
    ax.set_xticks(x)

    # Rotate subject names/labels
    rotation = 0
    ha = "center"
    if len(subjects) > 8:
        rotation = 45
        ha = "right"

    ax.set_xticklabels(subjects, fontsize=12, fontweight="bold", rotation=rotation, ha=ha)
    ax.set_ylim(0, 90)
    ax.tick_params(axis="y", direction="in", labelsize=11)
    ax.tick_params(axis="x", length=0)  # Hide x ticks marks

    # Legend at Top Horizontal
    ax.legend(
        loc="upper center",
        bbox_to_anchor=(0.5, 1.1),
        ncol=num_exams,
        frameon=False,
        fontsize=13,
        borderaxespad=0,
    )

    # Leave room for bottom labels(if rotated) and top legend
    plt.subplots_adjust(top=0.88, bottom=0.15)

    buf = io.BytesIO()
    plt.savefig(buf, format="png", dpi=100, bbox_inches="tight", transparent="True", pad_inches=0)
    plt.close()
    buf.seek(0)
    return buf


def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    tcVAlign = tcPr.find(f"{{{nsdecls('w')}}}shd")
    if tcVAlign is None:
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        tcPr.append(shd)
    else:
        tcVAlign.set(f"{{{nsdecls('w')}}}fill", color_hex)


def generate_doc(students, subjects, output_path, class_name):
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(10)

    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.4)
        section.right_margin = Inches(0.4)

    FULL_WIDTH_INCHES = 7.7  # A4 with 0.4 margins

    col_exam_w = 0.6
    remaining_w = FULL_WIDTH_INCHES - col_exam_w - 0.5
    if len(subjects) > 0:
        col_subj_w = remaining_w / len(subjects)
    else:
        col_subj_w = 0.5

    total_cols = 1 + len(subjects)

    total_students = len(students)

    for idx, student in enumerate(students):
        # Header Block
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Header Table
        header_table = doc.add_table(rows=4, cols=1)
        header_table.style = "Table Grid"
        header_table.autofit = False
        header_table.columns[0].width = Inches(FULL_WIDTH_INCHES)

        para._element.addnext(header_table._element)
        para._element.getparent().remove(para._element)

        # 1. Logo
        cell_img = header_table.rows[0].cells[0]
        p_img = cell_img.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if os.path.exists(HEADER_IMAGE):
            run = p_img.add_run()
            run.add_picture(HEADER_IMAGE, width=Inches(FULL_WIDTH_INCHES), height=Inches(2))

        # 2. School Name
        cell = header_table.rows[1].cells[0]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(TEXT_SCHOOL)
        run.bold = True
        run.font.size = Pt(14)

        # 3. Report Name
        cell = header_table.rows[2].cells[0]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(TEXT_REPORT)
        run.bold = True
        run.font.size = Pt(12)

        # 4. Class
        cell = header_table.rows[3].cells[0]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = p.add_run(class_name)
        run.bold = True
        run.font.size = Pt(12)

        doc.add_paragraph()  # Spacer

        # highlighted Student Info
        # Using a table with a background color
        info_table = doc.add_table(rows=1, cols=2)
        info_table.autofit = False
        per_cell_width = FULL_WIDTH_INCHES / 3
        info_table.rows[0].cells[0].width = Inches(per_cell_width)
        info_table.rows[0].cells[1].width = Inches(per_cell_width * 2)

        # Highlight color (Light Blue)
        HIGHLIGHT_COLOR = "CFE2F3"

        c0 = info_table.rows[0].cells[0]
        c0.text = "S. No.:"
        p = c0.paragraphs[0]
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(15)
        run.add_text(f"  {student['s_no']}")
        set_cell_background(c0, HIGHLIGHT_COLOR)

        c1 = info_table.rows[0].cells[1]
        c1.text = "Student Name:"
        p = c1.paragraphs[0]
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(15)
        run.add_text(f"  {student['name']}")
        set_cell_background(c1, HIGHLIGHT_COLOR)

        doc.add_paragraph()  # Spacer before marks table

        # Student marks table
        table = doc.add_table(rows=1, cols=total_cols)
        table.style = "Table Grid"
        table.autofit = False

        # Set widths
        table.columns[0].width = Inches(col_exam_w)
        for i in range(len(subjects)):
            table.columns[1 + i].width = Inches(col_subj_w)

        # Header Row
        hdr_cells = table.rows[0].cells
        headers = ["EXAM"] + subjects
        PINK_COLOR = "EAD1DC"

        for i, h in enumerate(headers):
            cell = hdr_cells[i]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.text = h
            set_cell_background(cell, PINK_COLOR)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)

        # Add rows for exams (e.g. 3 rows)
        exam_list = student["exams"]
        created_rows = []
        for _ in exam_list:
            row = table.add_row()
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height = Inches(0.3)
            created_rows.append(row)

        # Fill Data
        for row_idx, exam_data in enumerate(exam_list):
            cells = created_rows[row_idx].cells

            # Exam Name
            cells[0].text = exam_data["exam"]

            # Scores
            for subj_idx, subj in enumerate(subjects):
                val = exam_data["scores"].get(subj, 0.0)
                # Format: 1 decimal place?
                cells[1 + subj_idx].text = f"{val:.1f}"

            # Formatting
            for c in cells:
                c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                p = c.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.size = Pt(9)

        doc.add_paragraph()  # Spacer

        # Graph
        graph_img = create_graph(student, subjects)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(graph_img, width=Inches(FULL_WIDTH_INCHES))

        # Page Break
        if idx < total_students - 1:
            doc.add_page_break()

    doc.save(output_path)
    print(f"Document saved to {output_path}")


if __name__ == "__main__":
    if not os.path.exists(OUTPUT_DIR):
        os.makedirs(OUTPUT_DIR)

    if os.path.exists(INPUT_DIR):
        files = [f for f in os.listdir(INPUT_DIR) if f.casefold().endswith(".csv")]
        print(f"Found {len(files)} CSV files in {INPUT_DIR}")

        for filename in files:
            csv_path = os.path.join(INPUT_DIR, filename)

            # remove extension and use as class name
            base_name = os.path.splitext(filename)[0]
            class_name = f"CLASS {base_name.upper()}"

            # Try to make CLASS name nicer if it follows pattern "RESULT SHEET ... - Grade X"
            if " - " in base_name:
                parts = base_name.split(" - ")
                if len(parts) > 1:
                    class_name = f"CLASS {parts[-1].upper()}"

            output_filename = f"{base_name}_Report.docx"
            output_path = os.path.join(OUTPUT_DIR, output_filename)

            print(f"Processing {filename} -> {output_filename}...")

            students, subjects = parse_csv(csv_path)
            if students:
                print(f"  Found {len(students)} students and {len(subjects)} subjects.")
                try:
                    generate_doc(students, subjects, output_path, class_name)
                    print("  Done.")
                except Exception as e:
                    print(f"  Error generating doc: {e}")
            else:
                print("  No students found or failed to parse.")
    else:
        print(f"Input directory '{INPUT_DIR}' not found.")
